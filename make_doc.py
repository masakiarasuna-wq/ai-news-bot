from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ページ設定
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x58, 0x65, 0xF2)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.runs[0].font.size = Pt(11)
    return p

# タイトル
title = doc.add_heading('AIニュースBot 仕組み解説', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# これは何のためのシステム？
heading1('これは何のためのシステム？')
body('毎朝7時に、AI関連ニュースを自動で収集・要約して、Discordに通知するシステムです。PCがオフでもクラウド上で動くため、起動しておく必要はありません。')

doc.add_paragraph('')

# 各ツールの役割
heading1('登場するツールと役割')

heading2('1. Pythonスクリプト（ai_news_report.py）')
body('役割：システム全体の司令塔')
body('このファイルが「脳」にあたる部分です。「ニュースを取ってきて」「AIに要約させて」「Discordに送って」という一連の処理を順番に実行します。')

doc.add_paragraph('')

heading2('2. RSSフィード')
body('役割：ニュースの自動配信サービス')
body('ニュースサイトが提供している「新着記事の一覧データ」です。購読することで、サイトを直接開かなくても最新記事の情報（タイトル・URL・概要）を自動で取得できます。')
body('今回使っているソース：')
for s in [
    'TechCrunch AI（英語・グローバルIT/AI）',
    'The Verge AI（英語・消費者向けテクノロジー）',
    'VentureBeat AI（英語・ビジネス×AI）',
    'MIT Technology Review（英語・研究・テクノロジー）',
    'IEEE Spectrum（英語・ロボティクス・工学）',
    'ITmedia AI+（日本語・国内AI動向）',
    'Zenn AI（日本語・エンジニア向けAI技術）',
]:
    bullet(s)

doc.add_paragraph('')

heading2('3. Groq API')
body('役割：記事を日本語で要約するAI')
body('Groqというサービスが提供している無料のAI（LLama 3.3という大規模言語モデル）です。英語の記事を日本語に翻訳・要約したり、5つのカテゴリから重要記事を10件選んでくれます。無料枠が十分あり、登録するだけで使えます。')

doc.add_paragraph('')

heading2('4. Discord Bot')
body('役割：Discordに自動投稿するロボットアカウント')
body('Discord上で動くプログラム用のアカウントです。通常のユーザーと同じようにメッセージを送れますが、人間の代わりにプログラムが操作します。Botトークンは、BotがDiscordにログインするためのパスワードにあたります。')

doc.add_paragraph('')

heading2('5. 環境変数・シークレット')
body('役割：APIキーをファイルに直書きしないための仕組み')
body('DiscordトークンやGroq APIキーなどは「パスワード」と同じです。スクリプトのファイルに直接書いてしまうとセキュリティリスクがあります。そこで、OSやクラウドの「金庫」に別途保管しておき、実行時に読み込む仕組みを使っています。')
bullet('PC上：Windowsの「ユーザー環境変数」に保存')
bullet('GitHub Actions上：「リポジトリシークレット」に保存')

doc.add_paragraph('')

heading2('6. GitHub')
body('役割：コードの保存・管理場所')
body('プログラムのコードをインターネット上で保管するサービスです。今回はGitHub Actionsを使うためにコードを置く場所として利用しています。')

doc.add_paragraph('')

heading2('7. GitHub Actions')
body('役割：PCがオフでも自動実行してくれるクラウドサービス')
body('GitHubが提供する「クラウド上のスケジュール実行サービス」です。Pythonスクリプトをインターネット上のサーバーで動かしてくれるため、PCの電源が入っていなくても毎朝7時に自動実行されます。無料枠で十分使えます。')

doc.add_paragraph('')

# 全体の流れ
heading1('全体の流れ')
for step in [
    '① 毎朝7時になる',
    '② GitHub Actions がクラウド上でスクリプトを起動',
    '③ RSSフィードから各ニュースサイトの新着記事を収集',
    '④ Groq AI が重要な10記事を選んで日本語で要約',
    '⑤ Discord Bot がカード形式でDiscordに投稿',
    '⑥ スマホのDiscordで確認',
]:
    body(step)

doc.add_paragraph('')

# APIキー一覧（表）
heading1('使っているAPIキー一覧')
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'

headers = ['キー名', '用途', '料金']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

rows_data = [
    ('DISCORD_TOKEN', 'DiscordにBotとしてログインするため', '無料'),
    ('DISCORD_CHANNEL_ID', 'どのチャンネルに投稿するかの指定', '無料'),
    ('GROQ_API_KEY', 'Groq AIを使うための認証', '無料（一定量まで）'),
]
for i, (k, v, c) in enumerate(rows_data, 1):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    table.rows[i].cells[2].text = c

doc.save(r'C:\Users\masaki_arasuna\ai_news_bot\AIニュースbot_仕組み解説.docx')
print('完了')
